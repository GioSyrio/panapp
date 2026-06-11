#!/usr/bin/env python3
"""
build_questions_v2.py — DOCX-native question builder (Q+A)

Processes all 155 DOCX question + answer files and produces questions_v2.json
with structured sections and pre-rendered HTML.

Usage:
    python3 build_questions_v2.py
"""

import json, os, re, zipfile, io
from docx import Document
from PIL import Image

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data", "trapeza_data_1_3_218")
DOC_DIR = os.path.join(DATA_DIR, "attachments", "docx")
IMAGES_DIR = os.path.join(BASE_DIR, "static", "images", "exams")
QUESTIONS_FILE = os.path.join(DATA_DIR, "questions_classified.json")
OUTPUT_FILE = os.path.join(DATA_DIR, "questions_v2.json")
MAP_FILE = os.path.join(IMAGES_DIR, "diagram_map.json")

def e(s):
    return (s.replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))

def extract_images(docx_path, output_dir, qid):
    results = []
    os.makedirs(output_dir, exist_ok=True)
    try:
        with zipfile.ZipFile(docx_path,'r') as zf:
            for i,f in enumerate(sorted(f for f in zf.namelist() if f.startswith('word/media/'))):
                ext = os.path.splitext(f)[1].lower()
                data = zf.read(f)
                try:
                    img = Image.open(io.BytesIO(bytes(data))); w,h = img.size
                    if ext in ('.emf','.wmf') and (w<30 or h<30): continue
                    if ext in ('.emf','.wmf'): img = img.convert('RGB'); ext_out='.png'
                    else: ext_out = ext
                    if max(w,h) > 2000:
                        r = 2000/max(w,h); w,h = int(w*r), int(h*r)
                        img = img.resize((w,h), Image.LANCZOS)
                    name = f"{qid}_d{i+1}{ext_out}"; path = os.path.join(output_dir,name)
                    if name.endswith('.jpg') and img.mode in ('RGBA','P'): img = img.convert('RGB')
                    img.save(path)
                    results.append({"filename":name,"width":w,"height":h,
                                    "path":f"images/exams/{qid}/{name}"})
                except: pass
    except: pass
    return results

def is_code_line(t):
    if not t: return False
    if re.match(r'^(ΓΙΑ|ΟΣΟ|ΑΝ\b|ΑΛΛΙΩΣ|ΕΠΙΛΕΞΕ|ΕΠΑΝΑΛΑΒΕ|ΜΕΧΡΙΣ_ΟΤΟΥ|'
                r'ΔΙΑΒΑΣΕ|Διάβασε|ΓΡΑΨΕ|Γράψε|ΕΜΦΑΝΙΣΕ|Εμφάνισε|'
                r'ΚΑΛΕΣΕ|ΑΡΧΗ_ΕΠΑΝΑΛΗΨΗΣ|Αρχή_επανάληψης|'
                r'ΤΕΛΟΣ_ΕΠΑΝΑΛΗΨΗΣ|Τέλος_επανάληψης|'
                r'ΠΕΡΙΠΤΩΣΗ|Περίπτωση|ΤΕΛΟΣ_ΕΠΙΛΟΓΩΝ|Τέλος_επιλογών|'
                r'ΤΕΛΟΣ_ΑΝ|Τέλος_αν|ΤΕΛΟΣ\s+\w|'
                r'ΠΡΟΓΡΑΜΜΑ|ΜΕΤΑΒΛΗΤΕΣ|ΑΚΕΡΑΙΕΣ|ΠΡΑΓΜΑΤΙΚΕΣ|'
                r'ΧΑΡΑΚΤΗΡΕΣ|ΛΟΓΙΚΕΣ|ΣΤΑΘΕΡΕΣ|ΣΥΝΑΡΤΗΣΗ|'
                r'Αλγόριθμος|ΑΛΓΟΡΙΘΜΟΣ|ΔΙΑΔΙΚΑΣΙΑ)\b',t): return True
    if re.match(r'^[α-ωΑ-ΩίϊΐόύϋΰώάέήΆΈΉΊΌΎΏa-zA-Z0-9_]+\s*(←|<-|=)',t): return True
    if re.match(r'^\s{2,}(ΓΙΑ|ΟΣΟ|ΑΝ|ΔΙΑΒΑΣΕ|Διάβασε|ΓΡΑΨΕ|Γράψε|ΕΜΦΑΝΙΣΕ|Εμφάνισε)',t): return True
    return False

KW = (r'\b(ΓΙΑ|ΟΣΟ|ΑΝ\b|ΑΛΛΙΩΣ|ΕΠΙΛΕΞΕ|ΕΠΑΝΑΛΑΒΕ|ΜΕΧΡΙΣ_ΟΤΟΥ|'
      r'ΔΙΑΒΑΣΕ|Διάβασε|ΓΡΑΨΕ|Γράψε|ΕΜΦΑΝΙΣΕ|Εμφάνισε|'
      r'ΚΑΛΕΣΕ|ΑΡΧΗ_ΕΠΑΝΑΛΗΨΗΣ|Αρχή_επανάληψης|'
      r'ΤΕΛΟΣ_ΕΠΑΝΑΛΗΨΗΣ|Τέλος_επανάληψης|ΤΕΛΟΣ_ΑΝ|Τέλος_αν|'
      r'ΤΕΛΟΣ_ΕΠΙΛΟΓΩΝ|Τέλος_επιλογών|ΠΕΡΙΠΤΩΣΗ|Περίπτωση|'
      r'ΠΡΟΓΡΑΜΜΑ|ΜΕΤΑΒΛΗΤΕΣ|ΑΚΕΡΑΙΕΣ|ΠΡΑΓΜΑΤΙΚΕΣ|'
      r'ΧΑΡΑΚΤΗΡΕΣ|ΛΟΓΙΚΕΣ|ΣΤΑΘΕΡΕΣ|ΔΙΑΔΙΚΑΣΙΑ|ΣΥΝΑΡΤΗΣΗ|'
      r'Αλγόριθμος|ΑΛΓΟΡΙΘΜΟΣ)\b')

def norm(line):
    return re.sub(r'^\d{1,2}\s{1,4}', '', line)

def hl(line):
    line = norm(line)
    line = e(line)
    line = re.sub(KW, r'<span class="kw">\1</span>', line)
    line = re.sub(r'(?<!\w)(\d+(?:\.\d+)?)(?![\w;&])', r'<span class="num">\1</span>', line)
    line = re.sub(r'(←|<-)', r'<span class="op">\1</span>', line)
    return line

def _extract_cell_text_preserve_spaces(cell):
    xml = cell._element.xml
    paras = re.findall(r'<w:p[ >](.*?)</w:p>', xml, re.DOTALL)
    lines = []
    for p_xml in paras:
        texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', p_xml)
        lines.append(''.join(texts))
    return '\n'.join(lines)

# ── Question extraction ─────────────────────────────────────────────────────

def extract_q_sections(doc, qid, diagram_imgs):
    sections = []
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]; text = p.text.strip()
        if not text:
            for run in p.runs:
                d = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                b = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if d or b:
                    if diagram_imgs: sections.append({"type":"diagram","images":diagram_imgs})
            i += 1; continue

        if re.match(r'^ΘΕΜΑ\s+\d',text):
            sections.append({"type":"section_header","content":text}); i+=1; continue

        m = re.match(r'^(\d+\.\d+)[\.\s)]\s*(.*)',text)
        if m:
            subq = {"type":"sub_question","number":m.group(1),"content":m.group(2).strip()}
            i += 1
            while i < len(doc.paragraphs):
                nt = doc.paragraphs[i].text.strip()
                if not nt: i+=1; continue
                if re.match(r'^(ΘΕΜΑ|Μονάδες)\s',nt): break
                if re.match(r'^(\d+\.\d+)[\.\s)]',nt): break
                subq["content"] += " " + nt; i += 1
            sections.append(subq); continue

        m = re.match(r'Μονάδες\s+(\d+)',text)
        if m:
            sections.append({"type":"points","value":int(m.group(1))}); i+=1; continue

        m = re.match(r'^(Αλγόριθμος|ΑΛΓΟΡΙΘΜΟΣ)\s+(.+)',text)
        if m:
            algo = [text]; i += 1
            while i < len(doc.paragraphs):
                nt = doc.paragraphs[i].text; algo.append(nt)
                if re.match(r'^Τέλος\s+',nt.strip()): i+=1; break
                i += 1
            sections.append({"type":"algorithm_block","title":m.group(1)+" "+m.group(2),
                             "code_lines":[r.rstrip('\n') for r in algo]}); continue

        if is_code_line(text):
            lines = [text]; i += 1
            while i < len(doc.paragraphs):
                nt = doc.paragraphs[i].text; t2 = nt.strip()
                if not t2: lines.append(""); i+=1; continue
                if is_code_line(t2) or t2.startswith(" ") or re.match(r'^\d+\s',t2):
                    lines.append(nt.rstrip('\n')); i += 1
                else: break
            sections.append({"type":"code_block","code_lines":lines}); continue

        if re.search(r'\u2026\d+\u2026|\u2014\d+\u2014|_\w+_', text):
            gap_rows = [[text]]
            i += 1
            while i < len(doc.paragraphs):
                nt = doc.paragraphs[i].text.strip()
                if not nt: i += 1; continue
                if re.match(r'^(ΘΕΜΑ|Μονάδες)\s', nt): break
                if re.match(r'^(\d+\.\d+)[\.\s)]', nt): break
                if re.search(r'\u2026\d+\u2026|\u2014\d+\u2014|_\w+_', nt) or len(nt.split()) > 10:
                    gap_rows.append([nt]); i += 1
                else: break
            sections.append({"type":"gap_fill_table","rows":gap_rows})
            continue

        sections.append({"type":"text","content":text}); i+=1

    # Tables
    for table in doc.tables:
        rows = [[_extract_cell_text_preserve_spaces(c) for c in r.cells] for r in table.rows]
        if not rows: continue
        h = rows[0]
        is_match = any('Στήλη' in c for c in h) or any(any('Στήλη' in c for c in r) for r in rows[:3])
        has_gaps = any(re.search(r'\u2026\d+\u2026|\u2014\d+\u2014|_\w+_',c) for row in rows for c in row)
        has_code = False
        for row in rows:
            for cell in row:
                if '\n' in cell and len(cell)>50:
                    ls = cell.split('\n')
                    if sum(1 for l in ls if is_code_line(l.strip())) >= 3:
                        has_code = True; break
            if has_code: break
        if is_match:
            sections.append({"type":"matching_table","columns":[
                {"header":h[i] if i<len(h) else f"Col {i+1}",
                 "items":[r[i] if i<len(r) else "" for r in rows[1:]]}
                for i in range(len(h))]})
        elif has_gaps:
            sections.append({"type":"gap_fill_table","rows":rows})
        elif has_code:
            for row in rows:
                for cell in row:
                    if cell and '\n' in cell and len(cell)>50:
                        ls = cell.split('\n'); title = ""
                        if re.match(r'^(Αλγόριθμος|ΑΛΓΟΡΙΘΜΟΣ)\s+',ls[0].strip()):
                            title = ls[0].strip(); ls = ls[1:]
                        sections.append({"type":"algorithm_block","title":title,
                                         "code_lines":[l.rstrip('\r') for l in ls]})
        else:
            sections.append({"type":"table","rows":rows})
    return sections

# ── Answer extraction ───────────────────────────────────────────────────────

def extract_answer_steps(docx_path):
    if not os.path.exists(docx_path): return []
    try: doc = Document(docx_path)
    except: return []
    steps = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        m = re.match(r'^([A-ZΑ-Ω]\.|[\d]+\.|[\d]+\.[\d]+|[\u03B1-\u03C9]\)|[-–—])\s+(.*)', t)
        if m: steps.append({"label": m.group(1), "text": m.group(2)})
        elif steps and len(t) > 10: steps[-1]["text"] += " " + t
        elif len(t) > 10: steps.append({"label": "", "text": t})
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        code_lines = []
        for row in rows:
            for cell in row:
                if cell:
                    for line in cell.split('\n'): code_lines.append(line)
        if code_lines: steps.append({"label": "Code", "text": "", "code_lines": code_lines})
    return steps

def render_answer_html(ans_text, docx_path):
    steps = extract_answer_steps(docx_path)
    if not steps: return f'<div class="sol-text">{e(ans_text)}</div>'
    html = ''
    for i, step in enumerate(steps):
        label = step["label"] or f"{i+1}."
        html += '<div class="sol-step">'
        html += f'<div class="sol-step-label">{e(label)}</div>'
        if step.get("code_lines"):
            html += '<div class="code-body" style="margin-top:4px"><table>'
            for j, l in enumerate(step["code_lines"]):
                html += f'<tr><td class="ln">{j+1}</td><td class="lc">{hl(l)}</td></tr>'
            html += '</table></div>'
        else:
            html += f'<div class="sol-step-text">{e(step["text"])}</div>'
        html += '</div>'
    return html

# ── Question HTML ────────────────────────────────────────────────────────────

def render_q_html(sections):
    html = []
    for s in sections:
        t = s["type"]
        if t == "section_header":
            html.append(f'<div class="sec-header">{e(s["content"])}</div>')
        elif t == "sub_question":
            html.append(f'<div class="subq"><span class="subq-num">{e(s["number"])}</span> '
                        f'<span class="subq-text">{e(s["content"])}</span></div>')
        elif t == "points":
            html.append(f'<div class="points-chip">⭐ {s["value"]} μονάδες</div>')
        elif t in ("algorithm_block","code_block"):
            lines = s["code_lines"]; title = s.get("title","")
            html.append('<div class="code-card">')
            if title: html.append(f'<div class="code-card-header">🧮 {e(title)}</div>')
            html.append('<div class="code-body"><table>')
            for i2,l in enumerate(lines):
                html.append(f'<tr><td class="ln">{i2+1}</td><td class="lc">{hl(l)}</td></tr>')
            html.append('</table></div></div>')
        elif t == "text":
            html.append(f'<p class="text-content">{e(s["content"])}</p>')
        elif t == "diagram":
            imgs = s.get("images",[])
            html.append('<div class="diagram-gallery">')
            for img in imgs:
                html.append(f'<div class="diagram-item">'
                            f'<img src="/{img["path"]}" alt="Diagram" onclick="openDiagramModal(this.src)" loading="lazy">'
                            f'<div class="diagram-label">📐 {img["width"]}×{img["height"]}</div></div>')
            html.append('</div>')
        elif t == "matching_table":
            cols = s["columns"]; html.append('<table class="match-table"><thead><tr>')
            for c in cols: html.append(f'<th>{e(c["header"])}</th>')
            html.append('</tr></thead><tbody>')
            mx = max(len(c["items"]) for c in cols)
            for r in range(mx):
                html.append('<tr>')
                for c in cols:
                    v = e(c["items"][r]) if r < len(c["items"]) else ""
                    html.append(f'<td>{v}</td>')
                html.append('</tr>')
            html.append('</tbody></table>')
        elif t == "gap_fill_table":
            rows = s["rows"]; html.append('<table class="gap-table"><tbody>')
            for row in rows:
                html.append('<tr>')
                for cell in row:
                    highlighted = re.sub(r'(\u2026\d+\u2026|\u2014\d+\u2014|_\w+_)',
                                         r'<span class="gap-blank">\1</span>', e(cell))
                    html.append(f'<td>{highlighted}</td>')
                html.append('</tr>')
            html.append('</tbody></table>')
        elif t == "table":
            rows = s["rows"]; html.append('<table class="generic-table"><tbody>')
            for row in rows:
                html.append('<tr>' + ''.join(f'<td>{e(c)}</td>' for c in row) + '</tr>')
            html.append('</tbody></table>')
    return '\n'.join(html)

# ── Main ────────────────────────────────────────────────────────────────────

def main():
    if not os.path.exists(QUESTIONS_FILE):
        print(f"ERROR: {QUESTIONS_FILE} not found."); return
    with open(QUESTIONS_FILE, encoding="utf-8") as f: questions = json.load(f)
    dmap = {}
    if os.path.exists(MAP_FILE):
        with open(MAP_FILE, encoding="utf-8") as f: dmap = json.load(f)
    prev_v2 = {}
    if os.path.exists(OUTPUT_FILE):
        with open(OUTPUT_FILE, encoding="utf-8") as f:
            for p in json.load(f):
                prev_v2[str(p["id"])] = p
    output = []; processed = 0
    for q in questions:
        qid = q["id"]; qpath = os.path.join(DOC_DIR, f"{qid}-0.doc")
        apath = os.path.join(DOC_DIR, f"{qid}-4.doc")
        if not os.path.exists(qpath): continue
        print(f"  Q{qid}...", end=" ", flush=True)
        img_dir = os.path.join(IMAGES_DIR, str(qid))
        imgs = extract_images(qpath, img_dir, qid)
        all_imgs = list(imgs)
        if str(qid) in dmap: all_imgs = dmap[str(qid)].get("diagrams",[])
        doc_q = Document(qpath)
        sections = extract_q_sections(doc_q, qid, all_imgs)
        has_diag = any(s['type'] == 'diagram' for s in sections)
        if not has_diag and all_imgs:
            sections.append({"type": "diagram", "images": all_imgs})
        q_html = render_q_html(sections)
        ans_text = q.get("answer_text","")
        ans_html = render_answer_html(ans_text, apath)
        entry = prev_v2.get(str(qid), {})
        entry.update({
            "id": qid, "year": q.get("year"), "part": q.get("part"),
            "points": q.get("points"), "type": q.get("type","open_ended_problem"),
            "conceptual_tags": q.get("conceptual_tags",[]),
            "sections": sections, "question_html": q_html,
            "answer_html": ans_html, "answer_text": ans_text,
        })
        output.append(entry); processed += 1; print("✓")
    output.sort(key=lambda e: (-(e["year"] or 0), e["part"], e["id"]))
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\n  ✅ {processed} questions saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    main()