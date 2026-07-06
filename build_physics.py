#!/usr/bin/env python3
"""Build questions_v2.json for Physics from DOCX files — question_html + answer_html + part fix"""
import json, os, re, html as _html
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
DOC_DIR = os.path.join(BASE, "data/subjects/fysiki_prosanatolismoy/raw/docx")
V2_FILE = os.path.join(BASE, "data/subjects/fysiki_prosanatolismoy/questions_v2.json")

PART_MAP = {"ΘΕΜΑ 1": "Θέμα Α", "ΘΕΜΑ 2": "Θέμα Β", "ΘΕΜΑ 3": "Θέμα Γ", "ΘΕΜΑ 4": "Θέμα Δ",
            "ΘΕΜΑ Α": "Θέμα Α", "ΘΕΜΑ Β": "Θέμα Β", "ΘΕΜΑ Γ": "Θέμα Γ", "ΘΕΜΑ Δ": "Θέμα Δ"}

def esc(s):
    return _html.escape(s) if s else ""

def build_answer_html(docx_path):
    """Build answer_html from a physics answer DOCX file."""
    if not os.path.exists(docx_path):
        return ""
    doc = Document(docx_path)
    steps = []
    current_label = ""
    current_text = []
    
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        
        # Detect section header
        m = re.match(r'^ΘΕΜΑ\s+[Α-Δ1-4]', t)
        if m:
            continue  # Skip theme header in answer
        
        # Detect sub-question label: 2.1, 2.1.Α, 2.1.B, etc.
        m = re.match(r'^(\d+\.\d+(?:\.[Α-Ω])?)[\.\s)]*\s*(.*)', t)
        if m:
            # Save previous step
            if current_label or current_text:
                steps.append((current_label, " ".join(current_text)))
            current_label = m.group(1)
            current_text = [m.group(2)] if m.group(2) else []
            continue
        
        # Detect points line
        if re.match(r'^Μονάδες\s+\d+', t):
            if current_label or current_text:
                steps.append((current_label, " ".join(current_text)))
            current_label = ""
            current_text = []
            continue
        
        # Content line
        if current_label:
            current_text.append(t)
    
    # Save last step
    if current_label or current_text:
        steps.append((current_label, " ".join(current_text)))
    
    if not steps:
        return "<div class=\"sol-text\">Λύση στο αρχείο DOCX</div>"
    
    html = []
    for label, text in steps:
        if not text.strip():
            continue
        html.append('<div class="sol-step">')
        html.append(f'<div class="sol-step-label">{esc(label)}</div>')
        html.append(f'<div class="sol-step-text">{esc(text)}</div>')
        html.append('</div>')
    
    return "\n".join(html) if html else "<div class=\"sol-text\">Λύση στο αρχείο DOCX</div>"

def build_question_html(doc):
    """Build question_html from question DOCX."""
    html = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        
        if re.match(r'^ΘΕΜΑ\s+[Α-Δ1-4]', t):
            html.append(f'<div class="sec-header">{esc(t)}</div>')
        elif re.match(r'^Μονάδες\s+\d+', t):
            m = re.search(r'Μονάδες\s+(\d+)', t)
            html.append(f'<div class="points-chip">⭐ {m.group(1)} μονάδες</div>')
        elif re.match(r'^(\d+\.\d+(\.[Α-Ω])?|[α-ωΑ-Ω])[\s\)\.]', t):
            m = re.match(r'^((\d+\.\d+(?:\.[Α-Ω])?|[α-ωΑ-Ω]))[\s\)\.]*\s*(.*)', t)
            if m:
                num = m.group(1)
                text_part = m.group(3)
                html.append(f'<div class="subq"><span class="subq-num">{esc(num)})</span> <span class="subq-text">{esc(text_part)}</span></div>')
        else:
            html.append(f'<p class="text-content">{esc(t)}</p>')
    return "\n".join(html)

def detect_part_from_html(q_html):
    """Extract part from section header in question HTML."""
    m = re.search(r'ΘΕΜΑ\s+([Α-Δ1-4])', q_html)
    if m:
        key = f"ΘΕΜΑ {m.group(1)}"
        return PART_MAP.get(key, "Θέμα Α")
    return None

# ── Main ─────────────────────────────────────────────────────────────────────
v2 = json.load(open(V2_FILE, encoding="utf-8"))
rebuilt_q = 0
rebuilt_a = 0
fixed_parts = 0

for q in v2:
    qid = q["id"]
    qpath = os.path.join(DOC_DIR, f"{qid}-0.doc")
    apath = os.path.join(DOC_DIR, f"{qid}-4.doc")
    
    # ── Build question_html ──
    if os.path.exists(qpath):
        doc = Document(qpath)
        q_html = build_question_html(doc)
        if q_html:
            q["question_html"] = q_html
            rebuilt_q += 1
            
            # Detect and fix part label
            new_part = detect_part_from_html(q_html)
            if new_part and new_part != q.get("part"):
                q["part"] = new_part
                fixed_parts += 1
            
            # Ensure question_text
            if not q.get("question_text"):
                text = re.sub(r'<[^>]+>', ' ', q_html)
                q["question_text"] = re.sub(r'\s+', ' ', text).strip()[:3000]
    
    # ── Build answer_html ──
    if os.path.exists(apath):
        a_html = build_answer_html(apath)
        if a_html:
            q["answer_html"] = a_html
            rebuilt_a += 1
    
    # Ensure other fields
    try: q["year"] = int(q.get("year", 0))
    except: q["year"] = 2022
    if not q.get("points"):
        pts = re.findall(r'⭐ (\d+) μονάδες', q.get("question_html", ""))
        q["points"] = sum(int(p) for p in pts) if pts else 25

with open(V2_FILE, "w", encoding="utf-8") as f:
    json.dump(v2, f, ensure_ascii=False, indent=2)

# Summary
q0 = v2[0]
parts = {}
for q in v2:
    p = q.get("part", "?")
    parts[p] = parts.get(p, 0) + 1
empty_ans = sum(1 for q in v2 if not q.get("answer_html", "").strip())

print(f"✅ Question HTML rebuilt: {rebuilt_q}/{len(v2)}")
print(f"✅ Answer HTML rebuilt: {rebuilt_a}/{len(v2)}")
print(f"✅ Parts fixed: {fixed_parts}")
print(f"⚠️  Still empty answer_html: {empty_ans}/{len(v2)}")
print(f"📊 Parts distribution: {parts}")
print(f"📊 Hints intact: {bool(q0.get('hints'))}")