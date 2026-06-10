#!/usr/bin/env python3
"""
extract_doc_diagrams.py — Extracts diagrams from DOCX/DOC files

Sourced from the trapeza DOC files (manually saved by user as {qid}-0.doc).
These are ZIP-based Office files containing embedded images (JPEG/PNG/EMF)
and clean text with proper Greek encoding.

Advantages over PDF extraction:
  1. Native image resolution — not low-DPI screenshots
  2. Clean text extraction — no font encoding corruption
  3. Table structure preserved — matching tables parsed natively

Output: PNG files in static/images/exams/{question_id}/
         Updated mapping in static/images/exams/diagram_map.json

Usage:
    python3 extract_doc_diagrams.py                      # process all available DOCX
    python3 extract_doc_diagrams.py --id 25938           # single question
    python3 extract_doc_diagrams.py --dry-run            # preview only
"""

import json
import os
import sys
import argparse
import hashlib
import io
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image

# ── Paths ─────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data", "trapeza_data_1_3_218")
DOC_DIR = os.path.join(DATA_DIR, "attachments", "docx")
IMAGES_DIR = os.path.join(BASE_DIR, "static", "images", "exams")
MAP_FILE = os.path.join(IMAGES_DIR, "diagram_map.json")
QUESTIONS_FILE = os.path.join(DATA_DIR, "questions_classified.json")


def sanitize_filename(s):
    return s.replace("/", "_").replace("\\", "_").replace(":", "_").replace(" ", "_")


def find_doc_files():
    """Find all DOC files and return dict {qid: {question_path, answer_path}}."""
    result = {}
    if not os.path.exists(DOC_DIR):
        return result
    for fname in os.listdir(DOC_DIR):
        if not fname.endswith('.doc'):
            continue
        # Parse {qid}-{0|4}.doc
        base = os.path.splitext(fname)[0]
        parts = base.split('-')
        if len(parts) != 2:
            continue
        qid = parts[0]
        suffix = parts[1]
        if qid not in result:
            result[qid] = {}
        if suffix == '0':
            result[qid]['question'] = os.path.join(DOC_DIR, fname)
        elif suffix == '4':
            result[qid]['answer'] = os.path.join(DOC_DIR, fname)
    return result


def hash_file(path):
    """Generate content hash."""
    with open(path, "rb") as f:
        return hashlib.md5(f.read(16384)).hexdigest()[:8]


def extract_images_from_docx(docx_path, output_dir, question_id):
    """
    Extract all embedded images from a DOCX/DOC file.
    Reads all data upfront to avoid zipfile seek issues.
    Converts EMF/WMF to PNG using Pillow where possible.
    """
    results = []
    os.makedirs(output_dir, exist_ok=True)

    safe_qid = sanitize_filename(str(question_id))

    # First pass: read all media data (avoids zipfile seek issues)
    media_data = {}
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            for f in sorted(zf.namelist()):
                if f.startswith('word/media/'):
                    media_data[f] = bytes(zf.read(f))  # copy to RAM
    except Exception as e:
        print(f"  ❌ Failed to open {os.path.basename(docx_path)}: {e}")
        return results

    for i, (media_path, data) in enumerate(sorted(media_data.items())):
        ext = os.path.splitext(media_path)[1].lower()

        if ext in ('.emf', '.wmf'):
            out_name = f"{safe_qid}_diag{i + 1}.png"
        elif ext in ('.jpeg', '.jpg'):
            out_name = f"{safe_qid}_diag{i + 1}.jpg"
        else:
            out_name = f"{safe_qid}_diag{i + 1}{ext}"

        out_path = os.path.join(output_dir, out_name)

        if os.path.exists(out_path):
            try:
                with Image.open(out_path) as im:
                    w, h = im.size
                results.append({
                    "filename": out_name, "width": w, "height": h,
                    "format": ext.replace('.', ''),
                    "source": os.path.basename(media_path),
                    "path": f"images/exams/{safe_qid}/{out_name}",
                })
            except Exception:
                pass
            continue

        try:
            img = Image.open(io.BytesIO(data))
            w, h = img.size

            if ext in ('.emf', '.wmf'):
                if w < 30 or h < 30:
                    print(f"  ⏭️  {os.path.basename(media_path)} ({w}x{h}) — inline shape, skipping")
                    continue
                img = img.convert('RGB')
                max_dim = 2000
                if max(w, h) > max_dim:
                    ratio = max_dim / max(w, h)
                    w, h = int(w * ratio), int(h * ratio)
                    img = img.resize((w, h), Image.LANCZOS)
                img.save(out_path, 'PNG')
            else:
                if out_name.endswith('.jpg') and img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                img.save(out_path)

            results.append({
                "filename": out_name, "width": w, "height": h,
                "format": ext.replace('.', ''),
                "source": os.path.basename(media_path),
                "path": f"images/exams/{safe_qid}/{out_name}",
            })
            print(f"  🖼️  {os.path.basename(media_path)} → {out_name} ({w}x{h})")

        except Exception as e:
            err = str(e)
            if "WMF" in err and len(data) >= 8:
                magic = data[:4].hex()
                if magic == 'd7cdc69a':
                    print(f"  ⚠️  {os.path.basename(media_path)} is WMF (not EMF) — needs LibreOffice to convert")
                else:
                    print(f"  ⚠️  {os.path.basename(media_path)} failed: {err[:80]}")
            else:
                print(f"  ⚠️  {os.path.basename(media_path)} failed: {err[:80]}")

    return results


def extract_text_from_docx(docx_path):
    """Extract clean text from DOCX with paragraphs and tables."""
    try:
        doc = Document(docx_path)
        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Also extract table content
        for table in doc.tables:
            lines.append("")  # separator
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))

        return "\n".join(lines)
    except Exception as e:
        return f"[Error: {e}]"


def load_diagram_map():
    """Load existing diagram map."""
    if os.path.exists(MAP_FILE):
        with open(MAP_FILE, encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_diagram_map(dmap):
    """Save diagram map."""
    os.makedirs(os.path.dirname(MAP_FILE), exist_ok=True)
    with open(MAP_FILE, "w", encoding="utf-8") as f:
        json.dump(dmap, f, ensure_ascii=False, indent=2)


def main():
    parser = argparse.ArgumentParser(description="Extract diagrams from DOCX files")
    parser.add_argument("--id", type=int, default=0, help="Process single question by ID")
    parser.add_argument("--dry-run", action="store_true", help="Preview only, don't save")
    args = parser.parse_args()

    print("🔍 DOCX Diagram Extractor")
    print(f"   DOC source:  {DOC_DIR}")
    print(f"   Image output: {IMAGES_DIR}")
    print()

    # Find available DOC files
    doc_files = find_doc_files()
    if not doc_files:
        print("   ❌ No DOC files found. Run download_doc_files.py first.")
        sys.exit(1)

    print(f"   Found {len(doc_files)} questions with DOC files.")

    # Filter by --id
    if args.id:
        qid = str(args.id)
        if qid in doc_files:
            doc_files = {qid: doc_files[qid]}
        else:
            print(f"   ❌ No DOC file for question {args.id}")
            sys.exit(1)

    # Filter to questions that actually have given diagrams
    diagram_ids = set()
    if os.path.exists(QUESTIONS_FILE):
        with open(QUESTIONS_FILE, encoding="utf-8") as f:
            questions = json.load(f)
        for q in questions:
            qid = str(q["id"])
            # Check if question mentions diagram as GIVEN (not ask to create)
            qt = q.get("question_text", "")
            if any(kw in qt for kw in [
                "Δίνεται ο παρακάτω αλγόριθμος σε μορφή διαγράμματος",
                "Δίνεται το παρακάτω διάγραμμα ροής",
                "φαίνεται στο παρακάτω διάγραμμα",
                "ακόλουθο τμήμα διαγράμματος ροής να μετατραπεί",
            ]):
                diagram_ids.add(qid)

    # If --id not specified, only process diagram questions
    if not args.id:
        # Also include any question that has DOC files with embedded images
        for qid in list(doc_files.keys()):
            if qid in diagram_ids:
                continue
            # Check if DOC has images
            qpath = doc_files[qid].get('question')
            if qpath and os.path.exists(qpath):
                try:
                    with zipfile.ZipFile(qpath, 'r') as zf:
                        imgs = [f for f in zf.namelist() if f.startswith('word/media/')]
                        if imgs:
                            diagram_ids.add(qid)
                except Exception:
                    pass

    # Process each DOC file
    if not args.dry_run:
        diagram_map = load_diagram_map()
    else:
        diagram_map = {}

    total_images = 0
    processed = 0

    for qid in sorted(doc_files.keys(), key=int):
        if qid not in diagram_ids and not args.id:
            continue

        files = doc_files[qid]
        qpath = files.get('question')

        if not qpath or not os.path.exists(qpath):
            print(f"📄 Q{qid}: DOC missing")
            continue

        if args.dry_run:
            # Preview
            with zipfile.ZipFile(qpath, 'r') as zf:
                imgs = [f for f in zf.namelist() if f.startswith('word/media/')]
            print(f"📄 Q{qid}: {len(imgs)} embedded images")
            for img in imgs:
                info = zf.getinfo(img)
                print(f"    {os.path.basename(img)} ({info.file_size} bytes)")
            processed += 1
            continue

        print(f"📄 Q{qid} ({os.path.basename(qpath)})")
        output_dir = os.path.join(IMAGES_DIR, qid)
        results = extract_images_from_docx(qpath, output_dir, qid)

        if results:
            diagram_map[qid] = {
                "question_id": qid,
                "source": "docx",
                "content_hash": hash_file(qpath),
                "diagrams": results,
            }
            total_images += len(results)
        else:
            # Don't remove existing PDF entries if DOCX extraction produced nothing
            pass

        processed += 1

    if not args.dry_run:
        save_diagram_map(diagram_map)

    print(f"\n✅ Done! {processed} DOCX files processed.")
    print(f"   Diagram images: {total_images}")
    print(f"   Map file: {MAP_FILE}")

    if not args.dry_run and total_images > 0:
        print(f"\n   Updated diagram_map.json entries:")
        for qid in sorted(diagram_map.keys(), key=int):
            if diagram_map[qid].get("source") == "docx":
                count = len(diagram_map[qid].get("diagrams", []))
                print(f"     Q{qid}: {count} images")


if __name__ == "__main__":
    main()