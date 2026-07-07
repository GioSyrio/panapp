#!/usr/bin/env python3
"""Build questions_v2.json for Chemistry (Chimeia) from DOCX — answer_html + part fix + images"""
import json, os, re, html as _html, zipfile, io, argparse, shutil
from datetime import datetime
from docx import Document
from PIL import Image

BASE = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(BASE, "static", "images", "exams")
MAP_FILE = os.path.join(IMAGES_DIR, "diagram_map.json")

PART_MAP = {"ΘΕΜΑ 1": "Θέμα Α", "ΘΕΜΑ 2": "Θέμα Β", "ΘΕΜΑ 3": "Θέμα Γ", "ΘΕΜΑ 4": "Θέμα Δ",
            "ΘΕΜΑ Α": "Θέμα Α", "ΘΕΜΑ Β": "Θέμα Β", "ΘΕΜΑ Γ": "Θέμα Γ", "ΘΕΜΑ Δ": "Θέμα Δ"}
MAX_IMG_DIM = 1200

def esc(s):
    return _html.escape(s) if s else ""

def extract_images(docx_path, output_dir, qid):
    results = []
    os.makedirs(output_dir, exist_ok=True)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            for i, f in enumerate(sorted(f for f in zf.namelist() if f.startswith('word/media/'))):
                ext = os.path.splitext(f)[1].lower()
                data = zf.read(f)
                try:
                    img = Image.open(io.BytesIO(bytes(data)))
                    w, h = img.size
                    if w < 30 or h < 30: continue
                    if img.mode == 'P': img = img.convert('RGBA')
                    if img.mode == 'RGBA':
                        corners = [img.getpixel((0,0)), img.getpixel((w-1,0)), img.getpixel((0,h-1)), img.getpixel((w-1,h-1))]
                        if all(isinstance(c, tuple) and len(c)==4 and c[3]==0 for c in corners):
                            white_bg = Image.new('RGB', img.size, (255, 255, 255))
                            white_bg.paste(img, mask=img.split()[3])
                            img = white_bg
                    if ext in ('.emf', '.wmf'): img = img.convert('RGB'); ext_out = '.png'
                    else: ext_out = ext
                    if max(w, h) > MAX_IMG_DIM:
                        ratio = MAX_IMG_DIM / max(w, h)
                        w, h = int(w * ratio), int(h * ratio)
                        img = img.resize((w, h), Image.LANCZOS)
                    name = f"{qid}_d{i+1}{ext_out}"
                    path = os.path.join(output_dir, name)
                    if name.endswith('.jpg') and img.mode in ('RGBA', 'P'): img = img.convert('RGB')
                    img.save(path)
                    results.append({"filename": name, "width": w, "height": h, "path": f"images/exams/{qid}/{name}"})
                except: pass
    except: pass
    return results

def build_answer_html(docx_path):
    if not os.path.exists(docx_path): return ""
    doc = Document(docx_path)
    steps = []
    current_label, current_text = "", []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if re.match(r'^ΘΕΜΑ\s+[Α-Δ1-4]', t): continue
        m = re.match(r'^(\d+\.\d+(?:\.[Α-Ω])?)[\.\s)]*\s*(.*)', t)
        if m:
            if current_label or current_text: steps.append((current_label, " ".join(current_text)))
            current_label = m.group(1); current_text = [m.group(2)] if m.group(2) else []; continue
        if re.match(r'^Μονάδες\s+\d+', t):
            if current_label or current_text: steps.append((current_label, " ".join(current_text)))
            current_label = ""; current_text = []; continue
        if current_label: current_text.append(t)
    if current_label or current_text: steps.append((current_label, " ".join(current_text)))
    if not steps: return '<div class="sol-text">Λύση στο αρχείο DOCX</div>'
    html = []
    for label, text in steps:
        if not text.strip(): continue
        html.append('<div class="sol-step">')
        html.append(f'<div class="sol-step-label">{esc(label)}</div>')
        html.append(f'<div class="sol-step-text">{esc(text)}</div>')
        html.append('</div>')
    return "\n".join(html) if html else '<div class="sol-text">Λύση στο αρχείο DOCX</div>'

def build_question_html(doc, qid, images):
    html_parts, images_inserted = [], False
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            for run in p.runs:
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if (drawings or blips) and not images_inserted and images:
                    html_parts.append(_render_diagram_gallery(images)); images_inserted = True
            continue
        if re.match(r'^ΘΕΜΑ\s+[Α-Δ1-4]', t):
            html_parts.append(f'<div class="sec-header">{esc(t)}</div>')
            if not images_inserted and images: html_parts.append(_render_diagram_gallery(images)); images_inserted = True
        elif re.match(r'^Μονάδες\s+\d+', t):
            m = re.search(r'Μονάδες\s+(\d+)', t)
            html_parts.append(f'<div class="points-chip">⭐ {m.group(1)} μονάδες</div>')
        elif re.match(r'^(\d+\.\d+(\.[Α-Ω])?|[α-ωΑ-Ω])[\s\)\.]', t):
            m = re.match(r'^((\d+\.\d+(?:\.[Α-Ω])?|[α-ωΑ-Ω]))[\s\)\.]*\s*(.*)', t)
            if m:
                html_parts.append(f'<div class="subq"><span class="subq-num">{esc(m.group(1))})</span> <span class="subq-text">{esc(m.group(3))}</span></div>')
        else: html_parts.append(f'<p class="text-content">{esc(t)}</p>')
    if not images_inserted and images: html_parts.append(_render_diagram_gallery(images))
    return "\n".join(html_parts)

def _render_diagram_gallery(images):
    parts = ['<div class="diagram-gallery">']
    for img in images:
        parts.append(
            f'<div class="diagram-item">'
            f'<img src="/{img["path"]}" alt="Διάγραμμα" '
            f'style="max-width:100%;max-height:500px;border-radius:10px;'
            f'border:1px solid var(--border);cursor:pointer;'
            f'box-shadow:0 2px 8px rgba(0,0,0,0.06);" '
            f'onclick="openDiagramModal(this.src)" loading="lazy">'
            f'<div style="margin-top:6px;font-size:0.75rem;color:var(--text2);">'
            f'💡 Κάνε κλικ στο σχήμα για μεγέθυνση</div></div>')
    parts.append('</div>')
    return "\n".join(parts)

def detect_part_from_html(q_html):
    m = re.search(r'ΘΕΜΑ\s+([Α-Δ1-4])', q_html)
    if m: return PART_MAP.get(f"ΘΕΜΑ {m.group(1)}", "Θέμα Α")
    return None

def load_subject_config(subject_id):
    with open(os.path.join(BASE, "subjects", f"{subject_id}.json"), encoding="utf-8") as f:
        return json.load(f)

def main():
    p = argparse.ArgumentParser()
    p.add_argument("--subject", default="chimeia")
    args = p.parse_args()
    cfg = load_subject_config(args.subject)
    data_dir = os.path.join(BASE, cfg.get("data", {}).get("data_dir", f"data/subjects/{args.subject}"))
    doc_dir = os.path.join(data_dir, "raw", "docx")
    v2_file = os.path.join(data_dir, "questions_v2.json")
    if not os.path.exists(v2_file):
        print(f"ERROR: {v2_file} not found"); return
    backup = f"{v2_file}.bak.{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    shutil.copy2(v2_file, backup)
    print(f"📦 Backup: {os.path.basename(backup)}")
    v2 = json.load(open(v2_file, encoding="utf-8"))
    dmap = {}
    if os.path.exists(MAP_FILE):
        with open(MAP_FILE, encoding="utf-8") as f: dmap = json.load(f)
    rebuilt_q, rebuilt_a, fixed_parts, total_images = 0, 0, 0, 0
    for q in v2:
        qid = q["id"]
        qpath = os.path.join(doc_dir, f"{qid}-0.doc")
        apath = os.path.join(doc_dir, f"{qid}-4.doc")
        extracted_imgs = []
        if os.path.exists(qpath):
            img_dir = os.path.join(IMAGES_DIR, str(qid))
            if str(qid) in dmap:
                extracted_imgs = dmap[str(qid)].get("diagrams", [])
            else:
                extracted_imgs = extract_images(qpath, img_dir, qid)
                if extracted_imgs:
                    dmap[str(qid)] = {"diagrams": extracted_imgs, "source": "chimeia_docx"}
                    total_images += len(extracted_imgs)
        if os.path.exists(qpath):
            doc = Document(qpath)
            q_html = build_question_html(doc, qid, extracted_imgs)
            if q_html:
                q["question_html"] = q_html; rebuilt_q += 1
                new_part = detect_part_from_html(q_html)
                if new_part and new_part != q.get("part"):
                    q["part"] = new_part; fixed_parts += 1
                if not q.get("question_text"):
                    text = re.sub(r'<[^>]+>', ' ', q_html)
                    q["question_text"] = re.sub(r'\s+', ' ', text).strip()[:3000]
        if os.path.exists(apath):
            a_html = build_answer_html(apath)
            if a_html: q["answer_html"] = a_html; rebuilt_a += 1
        try: q["year"] = int(q.get("year", 0))
        except: q["year"] = 2022
        if not q.get("points"):
            pts = re.findall(r'⭐ (\d+) μονάδες', q.get("question_html", ""))
            q["points"] = sum(int(p) for p in pts) if pts else 25
    with open(v2_file, "w", encoding="utf-8") as f:
        json.dump(v2, f, ensure_ascii=False, indent=2)
    with open(MAP_FILE, "w", encoding="utf-8") as f:
        json.dump(dmap, f, ensure_ascii=False, indent=2)
    parts = {}
    for q in v2: parts[q.get("part", "?")] = parts.get(q.get("part", "?"), 0) + 1
    empty_ans = sum(1 for q in v2 if not q.get("answer_html", "").strip())
    print(f"🖼️  New images: {total_images} | Diagram map: {len(dmap)} questions")
    print(f"✅ QHTML: {rebuilt_q}/{len(v2)} | AHTML: {rebuilt_a}/{len(v2)} | Parts: {fixed_parts}")
    print(f"⚠️  Empty answers: {empty_ans}/{len(v2)} | Parts: {parts}")

if __name__ == "__main__":
    main()