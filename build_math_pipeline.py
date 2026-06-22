#!/usr/bin/env python3
"""
build_math_pipeline.py — DOCX-native math question builder

Processes 202 math DOCX question + answer files and produces questions_v2.json
with structured sections and pre-rendered HTML (KaTeX delimiters preserved).

Usage:
    python3 build_math_pipeline.py
    python3 build_math_pipeline.py --limit 5
"""

import json, os, re, sys, argparse
from docx import Document

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data", "subjects", "mathematics")
DOC_DIR = os.path.join(DATA_DIR, "raw", "docx")
QUESTIONS_FILE = os.path.join(DATA_DIR, "questions_classified.json")
OUTPUT_FILE = os.path.join(DATA_DIR, "questions_v2.json")

def e(s):
    return (s.replace("&", "&amp;").replace("<", "&lt;")
            .replace(">", "&gt;").replace('"', "&quot;"))

def extract_q_sections(doc):
    """Parse math question DOCX into structured sections."""
    sections = []
    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]; text = p.text.strip()
        if not text:
            i += 1; continue

        # ΘΕΜΑ header
        m = re.match(r'^ΘΕΜΑ\s+([Α-Δ])', text)
        if m:
            sections.append({"type": "section_header", "content": text}); i+=1; continue

        # Sub-question: α) or β) or γ) etc.
        m = re.match(r'^([α-ωΑ-Ω])\s*[\)\.]\s*(.*)', text)
        if m:
            subq = {"type": "sub_question", "number": m.group(1) + ")", "content": m.group(2).strip()}
            i += 1
            while i < len(doc.paragraphs):
                nt = doc.paragraphs[i].text.strip()
                if not nt: i+=1; continue
                if re.match(r'^(ΘΕΜΑ|Μονάδες)\s', nt): break
                if re.match(r'^([α-ωΑ-Ω])\s*[\)\.]\s*', nt): break
                subq["content"] += " " + nt; i += 1
            sections.append(subq); continue

        # Points
        m = re.match(r'Μονάδες\s+(\d+)', text)
        if m:
            sections.append({"type": "points", "value": int(m.group(1))}); i+=1; continue

        # Detect equations (preserve LaTeX)
        if re.search(r'[$\\]', text) and len(text) < 200:
            sections.append({"type": "equation", "content": text}); i+=1; continue

        # Regular text
        sections.append({"type": "text", "content": text}); i+=1

    # Tables
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if not rows: continue
        h = rows[0]
        is_match = any('Στήλη' in c for c in h)
        has_code = any(len(c) > 50 and '\n' in c for row in rows for c in row)
        if is_match:
            sections.append({"type": "matching_table", "columns": [
                {"header": h[i] if i < len(h) else f"Col {i+1}",
                 "items": [r[i] if i < len(r) else "" for r in rows[1:]]}
                for i in range(len(h))]})
        else:
            sections.append({"type": "table", "rows": rows})
    return sections

def extract_answer_sections(docx_path):
    """Parse math answer DOCX into step-by-step sections."""
    if not os.path.exists(docx_path): return [], ""
    try: doc = Document(docx_path)
    except: return [], ""
    
    steps = []
    full_text = ""
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        full_text += t + "\n"
        m = re.match(r'^([α-ωΑ-Ω])\s*[\)\.]\s*(.*)', t)
        if m:
            steps.append({"label": m.group(1) + ")", "content": m.group(2)})
        elif steps and len(t) > 10:
            steps[-1]["content"] += " " + t
        elif len(t) > 10:
            steps.append({"label": "", "content": t})
    
    # Check tables for answer code/steps
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        for row in rows:
            for cell in row:
                if cell:
                    steps.append({"label": "", "content": cell})
    
    return steps, full_text.strip()

def render_q_html(sections, qid=None):
    """Render math sections to HTML."""
    html = []
    for s in sections:
        t = s["type"]
        if t == "section_header":
            html.append(f'<div class="sec-header">{e(s["content"])}</div>')
        elif t == "sub_question":
            html.append(f'<div class="subq"><span class="subq-num">{e(s["number"])}</span> '
                        f'<span class="subq-text">{e(s["content"])}</span></div>')
        elif t == "points":
            html.append(f'<div class="points-chip">⭐ {s["value"]} μονάδες</div>')
        elif t == "equation":
            # Preserve LaTeX delimiters for KaTeX
            html.append(f'<div class="equation-display">{e(s["content"])}</div>')
        elif t == "text":
            content = e(s["content"])
            # Inject formula images (PNG/SVG only — skip WMF/EMF)
            formula_dir = os.path.join(BASE_DIR, "static", "images", "math_formulas", str(qid))
            if qid and os.path.isdir(formula_dir):
                for fname in sorted(os.listdir(formula_dir)):
                    if fname.lower().endswith(('.png', '.jpg', '.jpeg', '.svg')):
                        url = f"images/math_formulas/{qid}/{fname}"
                        content += f'<br><img src="{url}" class="formula-img" alt="formula" style="max-width:100%;margin:8px 0;">'
            html.append(f'<p class="text-content">{content}</p>')
        elif t == "matching_table":
            cols = s["columns"]; html.append('<table class="match-table"><thead><tr>')
            for c in cols: html.append(f'<th>{e(c["header"])}</th>')
            html.append('</tr></thead><tbody>')
            mx = max(len(c["items"]) for c in cols)
            for r in range(mx):
                html.append('<tr>')
                for c in cols:
                    v = e(c["items"][r]) if r < len(c["items"]) else ""
                    html.append(f'<td>{v}</td>')
                html.append('</tr>')
            html.append('</tbody></table>')
        elif t == "table":
            rows = s["rows"]; html.append('<table class="generic-table"><tbody>')
            for row in rows:
                html.append('<tr>' + ''.join(f'<td>{e(c)}</td>' for c in row) + '</tr>')
            html.append('</tbody></table>')
        elif t == "diagram":
            html.append('<div class="diagram-gallery">')
            for img in s.get("images", []):
                html.append(f'<img src="{img["path"]}" onclick="openDiagramModal(this.src)">')
            html.append('</div>')
    return '\n'.join(html)

def render_answer_html(steps):
    """Render answer sections to HTML."""
    html_parts = []
    for i, step in enumerate(steps):
        label = step["label"] or f"{i+1}."
        html_parts.append('<div class="sol-step">')
        html_parts.append(f'<div class="sol-step-label">{e(label)}</div>')
        html_parts.append(f'<div class="sol-step-text">{e(step["content"])}</div>')
        html_parts.append('</div>')
    return '\n'.join(html_parts)

def main():
    parser = argparse.ArgumentParser(description="Build math questions_v2.json from DOCX files")
    parser.add_argument("--limit", type=int, default=0, help="Process N questions only")
    args = parser.parse_args()

    if not os.path.exists(QUESTIONS_FILE):
        print(f"ERROR: {QUESTIONS_FILE} not found."); return

    with open(QUESTIONS_FILE, encoding="utf-8") as f:
        questions = json.load(f)

    # Load previous v2 to preserve extra fields (hints, etc.)
    prev_v2 = {}
    if os.path.exists(OUTPUT_FILE):
        with open(OUTPUT_FILE, encoding="utf-8") as f:
            for p in json.load(f):
                prev_v2[str(p["id"])] = p

    output = []; processed = 0
    total = args.limit or len(questions)
    for q in (questions[:args.limit] if args.limit else questions):
        qid = str(q["id"])
        qpath = os.path.join(DOC_DIR, f"{qid}-0.doc")
        apath = os.path.join(DOC_DIR, f"{qid}-4.doc")

        if not os.path.exists(qpath):
            print(f"  Q{qid}... skipping (no DOCX)")
            continue

        print(f"  Q{qid}...", end=" ", flush=True)

        try:
            doc_q = Document(qpath)
            sections = extract_q_sections(doc_q)
            q_html = render_q_html(sections, qid)

            ans_steps, ans_text = extract_answer_sections(apath)
            ans_html = render_answer_html(ans_steps) if ans_steps else f'<div class="sol-text">{e(ans_text)}</div>'

            entry = prev_v2.get(str(qid), {})
            entry.update({
                "id": q["id"], "year": q.get("year"), "part": q.get("part"),
                "points": q.get("points"), "type": q.get("type"),
                "conceptual_tags": q.get("conceptual_tags", []),
                "sections": sections, "question_html": q_html,
                "answer_html": ans_html, "answer_text": ans_text,
            })
            output.append(entry); processed += 1; print(f"✓ ({len(sections)} sections)")
        except Exception as e:
            print(f"❌ {e}")

    output.sort(key=lambda e: (str(e.get("year", "")), str(e.get("part", "")), e["id"]))
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\n  ✅ {processed} math questions saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    main()